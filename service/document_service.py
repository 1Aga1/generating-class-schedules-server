import io

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from models import Groups, Schedules, ScheduleParams, Subjects
from exceptions import ApiError


def generate_schedule(schedule_id: int):
    schedule = Schedules.get_or_none(Schedules.id == schedule_id)
    if not schedule:
        raise ApiError.BadRequest('Schedule not found')

    document = Document()
    document_name = 'Расписание на '+schedule.date.strftime('%d.%m.%Y')
    head = document.add_heading(document_name, level=2)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    section = document.sections[0]

    section.orientation = WD_ORIENTATION.LANDSCAPE

    section.page_width = Cm(27.94)
    section.page_height = Cm(21.59)

    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = document.styles['Normal']
    # название шрифта по умолчанию
    style.font.name = 'Times New Roman'
    # размер шрифта по умолчанию
    style.font.size = Pt(10)

    groups = [group for group in Groups.select().order_by(Groups.level.asc())]

    new_groups_list = []
    new_groups = []
    count = 0
    for group in groups:
        new_groups.append(group)
        count += 1
        if len(new_groups) == 6 or count == len(groups):
            new_groups_list.append(new_groups)
            new_groups = []

    for table_num, groups in enumerate(new_groups_list):
        document.add_paragraph()
        table = document.add_table(rows=9, cols=len(groups))
        table.style = 'Table Grid'
        document.add_section(start_type=WD_SECTION_START.ODD_PAGE)
        for column_num, group in enumerate(groups):
            cells = table.column_cells(column_num)
            for cell_id, cell in enumerate(cells):
                if cell_id == 0:
                    cell.text = group.name
                    rc = cell.paragraphs[0].runs[0]
                    rc.font.bold = True
                else:
                    params = ScheduleParams.select().where(ScheduleParams.schedule == schedule,
                                                           ScheduleParams.group == group,
                                                           ScheduleParams.number == cell_id)
                    for param in params:
                        subject = Subjects.get_or_none(id=param.subject)
                        cell.text += subject.name + ' - ' + subject.office + '\n'

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream, document_name
